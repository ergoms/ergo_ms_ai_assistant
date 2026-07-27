"""
Генератор Word документов.
Минимальная версия без сложных элементов.
"""
import logging
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)


class WordGenerator:
    """Генератор Word документов."""
    
    def generate_simple(
        self,
        output_path: Path,
        title: str,
        content: str,
        author: str = "AI",
        date: str = ""
    ) -> Path:
        """Генерирует простой документ."""
        
        doc = Document()
        
        doc.add_paragraph(self._safe(title))
        doc.add_paragraph("")
        
        if date:
            doc.add_paragraph(f"Дата: {date}")
        doc.add_paragraph(f"Автор: {author}")
        doc.add_paragraph("")
        
        if content:
            doc.add_paragraph(self._safe(content))
        
        doc.add_paragraph("")
        doc.add_paragraph("ERGO MS")
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return output_path
    
    def _safe(self, text) -> str:
        """Очищает текст."""
        if text is None:
            return ""
        
        text = str(text)
        
        # Удаляем управляющие символы
        result = []
        for char in text:
            code = ord(char)
            if code >= 32 or char in '\n\r\t':
                if not (0xD800 <= code <= 0xDFFF):
                    if code != 0xFFFE and code != 0xFFFF:
                        result.append(char)
        
        return ''.join(result)
