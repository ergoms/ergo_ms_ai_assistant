"""
Сервис для парсинга документов различных форматов (Word, PDF, TXT, MD, CSV, Excel).
"""
import logging
import os
import tempfile
from pathlib import Path
from typing import Optional, Tuple
from io import BytesIO

logger = logging.getLogger(__name__)

TABULAR_TYPES = frozenset({'csv', 'xlsx', 'xls'})
TEXT_LIKE_TYPES = frozenset({'txt', 'md', 'markdown'})


class DocumentParseError(Exception):
    """Общее исключение для ошибок парсинга документов."""
    pass


class DocumentParserService:
    """
    Сервис для извлечения текста из документов различных форматов
    
    Поддерживаемые форматы:
    - .docx (Microsoft Word)
    - .pdf (PDF документы)
    - .txt / .md (текстовые файлы)
    - .csv / .xlsx / .xls (таблицы через chonkie TableChef)
    """
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """
        Определяет тип файла по расширению
        
        Args:
            filename: Имя файла или путь
            
        Returns:
            Расширение файла в нижнем регистре (например, 'docx', 'pdf', 'txt')
        """
        return Path(filename).suffix.lower().lstrip('.') if Path(filename).suffix else ''
    
    @staticmethod
    def parse_document(file_path: Optional[str] = None, file_obj: Optional[BytesIO] = None, filename: Optional[str] = None) -> Tuple[str, str]:
        """
        Извлекает текст из документа
        
        Args:
            file_path: Путь к файлу на диске
            file_obj: Файловый объект (BytesIO)
            filename: Имя файла (для определения типа, если не указан путь)
            
        Returns:
            Кортеж (text_content, file_type):
            - text_content: Извлеченный текст
            - file_type: Тип файла (docx, pdf, txt и т.д.)
            
        Raises:
            DocumentParseError: При ошибке парсинга
        """
        # Определяем тип файла
        if file_path:
            file_type = DocumentParserService.get_file_type(file_path)
        elif filename:
            file_type = DocumentParserService.get_file_type(filename)
        else:
            raise DocumentParseError("Не указан ни путь к файлу, ни имя файла")
        
        if not file_type:
            raise DocumentParseError(f"Не удалось определить тип файла: {filename or file_path}")
        
        try:
            if file_type == 'docx':
                return DocumentParserService._parse_docx(file_path, file_obj), file_type
            if file_type == 'pdf':
                return DocumentParserService._parse_pdf(file_path, file_obj), file_type
            if file_type in TEXT_LIKE_TYPES:
                # Нормализуем markdown → md
                normalized = 'md' if file_type in ('md', 'markdown') else 'txt'
                return DocumentParserService._parse_txt(file_path, file_obj), normalized
            if file_type in TABULAR_TYPES:
                return DocumentParserService._parse_tabular(
                    file_path, file_obj, file_type,
                ), file_type
            raise DocumentParseError(f"Неподдерживаемый тип файла: {file_type}")
                
        except DocumentParseError:
            raise
        except Exception as e:
            raise DocumentParseError(f"Ошибка парсинга файла {file_type}: {str(e)}") from e
    
    @staticmethod
    def _parse_docx(file_path: Optional[str] = None, file_obj: Optional[BytesIO] = None) -> str:
        """
        Парсит DOCX файл
        
        Args:
            file_path: Путь к файлу
            file_obj: Файловый объект
            
        Returns:
            Извлеченный текст
        """
        try:
            from docx import Document
        except ImportError:
            raise DocumentParseError(
                "Модуль python-docx не установлен. "
                "Установите: poetry add python-docx"
            )
        
        try:
            if file_obj:
                # Используем файловый объект
                file_obj.seek(0)  # Перемещаемся в начало
                doc = Document(file_obj)
            elif file_path:
                # Читаем с диска
                doc = Document(file_path)
            else:
                raise DocumentParseError("Не указан ни путь к файлу, ни файловый объект")
            
            # Извлекаем текст из всех параграфов
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            result = "\n\n".join(text_parts)
            
            if not result.strip():
                raise DocumentParseError("Документ пуст или не удалось извлечь текст")
            
            return result.strip()
            
        except Exception as e:
            if isinstance(e, DocumentParseError):
                raise
            raise DocumentParseError(f"Ошибка парсинга DOCX: {str(e)}") from e
    
    @staticmethod
    def _parse_pdf(file_path: Optional[str] = None, file_obj: Optional[BytesIO] = None) -> str:
        """
        Парсит PDF файл
        
        Args:
            file_path: Путь к файлу
            file_obj: Файловый объект
            
        Returns:
            Извлеченный текст
        """
        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                import pdfplumber
                use_pdfplumber = True
            except ImportError:
                raise DocumentParseError(
                    "Модуль для парсинга PDF не установлен. "
                    "Установите один из: pypdf или pdfplumber"
                )
        else:
            use_pdfplumber = False
        
        try:
            text_parts = []
            
            if use_pdfplumber:
                # Используем pdfplumber (более точный парсинг)
                import pdfplumber
                
                if file_obj:
                    file_obj.seek(0)
                    pdf = pdfplumber.open(file_obj)
                elif file_path:
                    pdf = pdfplumber.open(file_path)
                else:
                    raise DocumentParseError("Не указан ни путь к файлу, ни файловый объект")
                
                try:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text.strip())
                finally:
                    pdf.close()
            else:
                # Используем pypdf
                if file_obj:
                    file_obj.seek(0)
                    pdf_reader = PdfReader(file_obj)
                elif file_path:
                    with open(file_path, 'rb') as f:
                        pdf_reader = PdfReader(f)
                else:
                    raise DocumentParseError("Не указан ни путь к файлу, ни файловый объект")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(page_text.strip())
                    except Exception as e:
                        logger.warning(f"Ошибка извлечения текста со страницы {page_num + 1}: {e}")
                        continue
            
            result = "\n\n".join(text_parts)
            
            if not result.strip():
                raise DocumentParseError("PDF документ пуст или не удалось извлечь текст")
            
            return result.strip()
            
        except Exception as e:
            if isinstance(e, DocumentParseError):
                raise
            raise DocumentParseError(f"Ошибка парсинга PDF: {str(e)}") from e
    
    @staticmethod
    def _parse_txt(file_path: Optional[str] = None, file_obj: Optional[BytesIO] = None) -> str:
        """
        Парсит текстовый файл
        
        Args:
            file_path: Путь к файлу
            file_obj: Файловый объект
            
        Returns:
            Содержимое файла
        """
        try:
            if file_obj:
                file_obj.seek(0)
                # Пробуем разные кодировки
                content = file_obj.read()
                for encoding in ['utf-8', 'utf-8-sig', 'cp1251', 'latin-1']:
                    try:
                        return content.decode(encoding)
                    except UnicodeDecodeError:
                        continue
                # Если не удалось декодировать, возвращаем с ошибками замены
                return content.decode('utf-8', errors='replace')
            elif file_path:
                # Пробуем разные кодировки
                for encoding in ['utf-8', 'utf-8-sig', 'cp1251', 'latin-1']:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            return f.read()
                    except UnicodeDecodeError:
                        continue
                # Если не удалось декодировать, возвращаем с ошибками замены
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    return f.read()
            else:
                raise DocumentParseError("Не указан ни путь к файлу, ни файловый объект")
                
        except Exception as e:
            if isinstance(e, DocumentParseError):
                raise
            raise DocumentParseError(f"Ошибка чтения текстового файла: {str(e)}") from e

    @staticmethod
    def _parse_tabular(
        file_path: Optional[str] = None,
        file_obj: Optional[BytesIO] = None,
        file_type: str = 'csv',
    ) -> str:
        """CSV/Excel → markdown через chonkie TableChef (локальный path)."""
        try:
            from chonkie import TableChef
        except ImportError as exc:
            raise DocumentParseError(
                'Модуль chonkie не установлен. Выполните: ergoms python-install'
            ) from exc

        temp_path: Optional[str] = None
        try:
            if file_path:
                path_for_chef = file_path
            elif file_obj:
                file_obj.seek(0)
                suffix = f'.{file_type}' if file_type else '.csv'
                fd, temp_path = tempfile.mkstemp(suffix=suffix)
                try:
                    with os.fdopen(fd, 'wb') as tmp:
                        tmp.write(file_obj.read())
                except Exception:
                    os.close(fd)
                    raise
                path_for_chef = temp_path
            else:
                raise DocumentParseError('Не указан ни путь к файлу, ни файловый объект')

            # Excel: добавляем имена листов для RAG (TableChef их не включает).
            if file_type in ('xlsx', 'xls'):
                return DocumentParserService._excel_to_markdown(path_for_chef)

            chef = TableChef()
            doc = chef.process(path_for_chef)
            content = (getattr(doc, 'content', None) or '').strip()
            if not content:
                raise DocumentParseError('Табличный файл пуст или не удалось извлечь данные')
            return content
        except DocumentParseError:
            raise
        except Exception as exc:
            raise DocumentParseError(f'Ошибка парсинга таблицы ({file_type}): {exc}') from exc
        finally:
            if temp_path and os.path.isfile(temp_path):
                try:
                    os.unlink(temp_path)
                except OSError:
                    pass

    @staticmethod
    def _excel_to_markdown(path: str) -> str:
        """Все листы Excel → markdown с заголовками листов."""
        try:
            import pandas as pd
        except ImportError as exc:
            raise DocumentParseError(
                'Pandas требуется для Excel. Выполните: ergoms python-install'
            ) from exc

        try:
            sheets = pd.read_excel(path, sheet_name=None)
        except Exception as exc:
            raise DocumentParseError(f'Ошибка чтения Excel: {exc}') from exc

        if not sheets:
            raise DocumentParseError('Excel-файл не содержит листов')

        parts = []
        for sheet_name, frame in sheets.items():
            try:
                table_md = frame.to_markdown(index=False) or ''
            except Exception:
                table_md = frame.to_string(index=False)
            parts.append(f'## {sheet_name}\n\n{table_md}'.strip())

        result = '\n\n'.join(parts).strip()
        if not result:
            raise DocumentParseError('Excel документ пуст или не удалось извлечь данные')
        return result
